"""Extract plain text from PRD sources: DOCX, PDF, ODT."""

from __future__ import annotations

import io
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path


def extract_from_bytes(filename: str, data: bytes) -> str:
    """Route by extension and return UTF-8 text."""
    ext = Path(filename).suffix.lower()
    if ext == ".docx":
        return _docx_text(data)
    if ext == ".pdf":
        return _pdf_text(data)
    if ext == ".odt":
        return _odt_text(data)
    raise ValueError(f"Unsupported file type: {ext}. Use .docx, .pdf, or .odt.")


def _docx_text(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts) if parts else ""


def _pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text()
        if t and t.strip():
            parts.append(t.strip())
    return "\n\n".join(parts) if parts else ""


_TEXT_NS = "urn:oasis:names:tc:opendocument:xmlns:text:1.0"
_P = f"{{{_TEXT_NS}}}p"
_H = f"{{{_TEXT_NS}}}h"


def _odt_text(data: bytes) -> str:
    """ODT is a ZIP with content.xml; extract text from text:p and text:h."""
    buf = io.BytesIO(data)
    if not zipfile.is_zipfile(buf):
        raise ValueError("Invalid ODT file (not a ZIP archive).")
    with zipfile.ZipFile(buf, "r") as zf:
        try:
            xml_bytes = zf.read("content.xml")
        except KeyError as e:
            raise ValueError("Invalid ODT: missing content.xml") from e

    root = ET.fromstring(xml_bytes)
    chunks: list[str] = []
    for el in root.iter():
        if el.tag in (_P, _H):
            text = "".join(el.itertext()).strip()
            if text:
                chunks.append(text)
    return "\n\n".join(chunks) if chunks else ""
