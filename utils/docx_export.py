"""Convert markdown-ish story output to a formatted Word document (.docx)."""

from __future__ import annotations

import io
import re

_CHECKBOX_RE = re.compile(r"^(\s*)-\s*\[[ xX]\]\s+")
_BULLET_RE = re.compile(r"^(\s*)[-*]\s+")
_NUMBER_RE = re.compile(r"^(\s*)(\d+)\.\s+")


def _add_italic_segments(paragraph, text: str) -> None:
    """Add *italic* segments; ** is not expected here."""
    pos = 0
    for m in re.finditer(r"(?<!\*)\*(?!\*)([^*]+)\*(?!\*)", text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        paragraph.add_run(m.group(1)).italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_runs(paragraph, text: str) -> None:
    """**bold** and *italic* (italic does not nest inside bold in this MVP)."""
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if not part:
            continue
        if len(part) >= 4 and part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            _add_italic_segments(paragraph, part)


def _strip_heading(line: str) -> tuple[int, str] | None:
    m = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
    if not m:
        return None
    level = min(len(m.group(1)), 9)
    return level, m.group(2).strip()


def _list_payload(line: str) -> tuple[str, str] | None:
    """
    Returns (kind, body) where kind is 'bullet', 'number', or 'checkbox'.
    """
    m = _CHECKBOX_RE.match(line)
    if m:
        return "bullet", line[m.end() :].strip()
    m = _NUMBER_RE.match(line)
    if m:
        return "number", line[m.end() :].strip()
    m = _BULLET_RE.match(line)
    if m:
        return "bullet", line[m.end() :].strip()
    return None


def markdown_to_docx_bytes(markdown: str, title: str = "User stories") -> bytes:
    """Build a .docx from markdown text; returns file bytes."""
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(6)

    doc.core_properties.title = title
    doc.add_heading(title, level=0)

    number_fallback = 0

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        if stripped in ("---", "***", "___"):
            doc.add_paragraph().paragraph_format.space_before = Pt(4)
            continue

        hd = _strip_heading(stripped)
        if hd:
            level, text = hd
            doc.add_heading(text, level=level)
            continue

        listed = _list_payload(line)
        if listed is not None:
            kind, body = listed
            style = "List Number" if kind == "number" else "List Bullet"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph()
                if kind == "number":
                    number_fallback += 1
                    p.add_run(f"{number_fallback}. ")
                else:
                    p.add_run("• ")
            _add_runs(p, body)
            continue

        p = doc.add_paragraph()
        _add_runs(p, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
